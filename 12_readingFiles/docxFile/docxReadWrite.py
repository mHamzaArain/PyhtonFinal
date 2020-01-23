import docx

# d = docx.Document(<.docx>)                # To read file
# d.paragraphs                                           # All paragraphs as objs 
# d.paragraphs[0]                                       # 1 para as obj at 0 index
# d.paragraphs[0].text                                # Para as str at 1st index 
# p.runs[1].text                                            # bold, underline, italic as str
# p.runs[1].bold                                           # bold as obj                
# p.runs[3].italic = True                               # To make font italic
# p.runs[3]..text= 'italic to underline'           # To change text 
# p.style                                                        #  str style as obj
# p.style = 'Title'                                           # To change style 


# ##############Implementations
# d = docx.Document(r'C:\Users\Family\Desktop\demo.docx')
# print(d.paragraphs)     # [<docx.text.paragraph.Paragraph object at 0x000000000360ED68>, <docx.text.paragraph.Paragraph object at 0x000000000360ECC0>, <docx.text.paragraph.Paragraph object at 0x000000000360EDD8>, <docx.text.paragraph.Paragraph object at 0x000000000360ED30>, <docx.text.paragraph.Paragraph object at 0x000000000360EDA0>, <docx.text.paragraph.Paragraph object at 0x000000000360EE10>, <docx.text.paragraph.Paragraph object at 0x000000000360EE48>]
# d.paragraphs[0]     # <docx.text.paragraph.Paragraph object at 0x000000000360ECF8>
# print(d.paragraphs[0].text)     # 'Document Title'
# p.runs[0].text
# 'A plain paragraph having some '
# p.runs[1].text
# 'bold'
# p.runs[3].text
# 'italic.'
# p.runs[3].bold
# p.runs[3].bold = None
# p.runs[1].text = None

# p.runs[1].bold = None
# p.runs[0].bold = None
# p.runs[0].bold = None
# p.runs[0].bold
# p.runs[3].italic
# True
# p.runs[3].italic = True
# p.runs[3].text = 'italic to underline'
# p.style
# # _ParagraphStyle('Normal') id: 56683992
# p.style = 'Title'
# d = docx.Document()
# d.add_paragraph('This is paragraph')
# # <docx.text.paragraph.Paragraph object at 0x0000000003631EF0>
# d.add_paragraph('This is another paragraph')
# # <docx.text.paragraph.Paragraph object at 0x000000000360EDD8>
# p.runs        # [<docx.text.run.Run object at 0x0000000003631EB8>]
# p.runs[1].bold = True
# ##############Read File
# def getText(filename):
#     doc = docx.Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     return '\n'.join(fullText)

# print(getText(r"C:\Users\Family\Desktop\demo.docx"))
